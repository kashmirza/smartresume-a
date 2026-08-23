"""
PDF & DOCX Document Generation Engine for SmartResume AI.

Uses ReportLab to generate ATS-friendly PDFs across multiple templates
(ats_classic, modern, minimal) and python-docx to generate DOCX files.
"""

from typing import Dict, List, Optional, Any, Union
import io

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, KeepTogether
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class PDFGenerator:
    """Production PDF and DOCX document generator for resumes."""

    def __init__(self):
        """Initialize PDF generator."""
        pass

    def generate_pdf(
        self,
        resume_data: Union[Dict[str, Any], Any],
        template_name: str = "ats_classic"
    ) -> bytes:
        """
        Generate ATS-compliant PDF document from resume data.

        Args:
            resume_data: Structured resume dictionary or object.
            template_name: Template identifier ('ats_classic', 'modern', 'minimal').

        Returns:
            PDF file content in bytes.
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab library is required for PDF generation.")

        data = self._normalize_resume_data(resume_data)
        buffer = io.BytesIO()

        template = template_name.lower().strip()
        if template not in ["ats_classic", "modern", "minimal"]:
            template = "ats_classic"

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )

        styles = self._build_styles(template)
        story = []

        # 1. Header (Name + Contact Info)
        self._build_header(data, template, styles, story)

        # 2. Professional Summary
        summary = data.get("summary") or data.get("objective")
        if summary:
            self._add_section_header("PROFESSIONAL SUMMARY", template, styles, story)
            story.append(Paragraph(str(summary), styles["Summary"]))
            story.append(Spacer(1, 10))

        # 3. Work Experience
        experience = data.get("experience", []) or data.get("work_experience", [])
        if experience:
            self._add_section_header("WORK EXPERIENCE", template, styles, story)
            for exp in experience:
                if isinstance(exp, dict):
                    title = exp.get("position") or exp.get("title", "Position")
                    company = exp.get("company", "")
                    dates = exp.get("dates") or f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                    location = exp.get("location", "")

                    header_text = f"<b>{title}</b> — {company}"
                    if location:
                        header_text += f" ({location})"

                    story.append(Paragraph(
                        f"<font color='{styles['accent_hex']}'>{header_text}</font>",
                        styles["ItemHeader"]
                    ))
                    if dates.strip(" -"):
                        story.append(Paragraph(f"<i>{dates}</i>", styles["ItemMeta"]))

                    bullets = exp.get("highlights") or exp.get("bullets") or exp.get("description", [])
                    if isinstance(bullets, str):
                        bullets = [bullets]

                    for bullet in bullets:
                        if bullet and str(bullet).strip():
                            story.append(Paragraph(f"• {str(bullet).strip()}", styles["Bullet"]))

                    story.append(Spacer(1, 8))

        # 4. Technical Skills
        skills = data.get("skills", [])
        if skills:
            self._add_section_header("TECHNICAL SKILLS", template, styles, story)
            if isinstance(skills, list):
                skills_str = ", ".join([str(s) for s in skills if s])
                story.append(Paragraph(skills_str, styles["BodyText"]))
            elif isinstance(skills, dict):
                for category, cat_skills in skills.items():
                    if isinstance(cat_skills, list) and cat_skills:
                        c_str = ", ".join([str(s) for s in cat_skills if s])
                        cat_title = category.replace("_", " ").title()
                        story.append(Paragraph(f"<b>{cat_title}:</b> {c_str}", styles["BodyText"]))
            story.append(Spacer(1, 10))

        # 5. Key Projects
        projects = data.get("projects", [])
        if projects:
            self._add_section_header("KEY PROJECTS", template, styles, story)
            for proj in projects:
                if isinstance(proj, dict):
                    name = proj.get("name") or proj.get("title", "Project")
                    tech = proj.get("tech_stack") or proj.get("technologies")
                    desc = proj.get("description", "")

                    p_header = f"<b>{name}</b>"
                    if tech:
                        tech_str = ", ".join(tech) if isinstance(tech, list) else str(tech)
                        p_header += f" | <i>{tech_str}</i>"

                    story.append(Paragraph(p_header, styles["ItemHeader"]))
                    if desc:
                        story.append(Paragraph(str(desc), styles["BodyText"]))

                    bullets = proj.get("highlights") or proj.get("bullets", [])
                    if isinstance(bullets, list):
                        for b in bullets:
                            story.append(Paragraph(f"• {str(b)}", styles["Bullet"]))

                    story.append(Spacer(1, 6))

        # 6. Education
        education = data.get("education", [])
        if education:
            self._add_section_header("EDUCATION", template, styles, story)
            for edu in education:
                if isinstance(edu, dict):
                    degree = edu.get("degree", "Degree")
                    institution = edu.get("institution") or edu.get("school", "")
                    dates = edu.get("dates") or edu.get("year", "")
                    gpa = edu.get("gpa")

                    edu_line = f"<b>{degree}</b> — {institution}"
                    if dates:
                        edu_line += f" ({dates})"
                    if gpa:
                        edu_line += f" | GPA: {gpa}"

                    story.append(Paragraph(edu_line, styles["BodyText"]))
            story.append(Spacer(1, 10))

        # 7. Certifications
        certifications = data.get("certifications", [])
        if certifications:
            self._add_section_header("CERTIFICATIONS", template, styles, story)
            for cert in certifications:
                if isinstance(cert, dict):
                    cert_name = cert.get("name") or cert.get("title", "")
                    issuer = cert.get("issuer") or cert.get("authority", "")
                    date = cert.get("date") or cert.get("year", "")
                    line = f"• <b>{cert_name}</b>"
                    if issuer:
                        line += f" — {issuer}"
                    if date:
                        line += f" ({date})"
                    story.append(Paragraph(line, styles["BodyText"]))
                elif isinstance(cert, str):
                    story.append(Paragraph(f"• {cert}", styles["BodyText"]))

        # Build PDF
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        return pdf_data

    def generate_docx(
        self,
        resume_data: Union[Dict[str, Any], Any],
        template_name: str = "ats_classic"
    ) -> bytes:
        """
        Generate DOCX file from resume data using python-docx.

        Returns bytes representing the DOCX file stream.
        """
        if not DOCX_AVAILABLE:
            note = f"DOCX generation unavailable: python-docx not installed. Template: {template_name}"
            return note.encode("utf-8")

        data = self._normalize_resume_data(resume_data)
        doc = docx.Document()

        # Page Margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        contact = data.get("contact", {}) or data.get("contact_info", {})
        name = contact.get("name", "Applicant Name")

        # Name
        p_name = doc.add_paragraph()
        run_name = p_name.add_run(name)
        run_name.font.size = Pt(20)
        run_name.font.bold = True
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact Info Line
        c_parts = []
        for field in ["email", "phone", "location", "linkedin", "github"]:
            val = contact.get(field)
            if val:
                c_parts.append(str(val))

        if c_parts:
            p_contact = doc.add_paragraph()
            run_c = p_contact.add_run(" | ".join(c_parts))
            run_c.font.size = Pt(9)
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        summary = data.get("summary") or data.get("objective")
        if summary:
            doc.add_heading("PROFESSIONAL SUMMARY", level=2)
            doc.add_paragraph(str(summary))

        # Experience
        experience = data.get("experience", []) or data.get("work_experience", [])
        if experience:
            doc.add_heading("WORK EXPERIENCE", level=2)
            for exp in experience:
                if isinstance(exp, dict):
                    title = exp.get("position") or exp.get("title", "Position")
                    company = exp.get("company", "")
                    dates = exp.get("dates", "")
                    p_exp = doc.add_paragraph()
                    r_title = p_exp.add_run(f"{title} - {company} ")
                    r_title.bold = True
                    if dates:
                        r_date = p_exp.add_run(f"({dates})")
                        r_date.italic = True

                    bullets = exp.get("highlights") or exp.get("bullets", [])
                    if isinstance(bullets, str):
                        bullets = [bullets]
                    for b in bullets:
                        doc.add_paragraph(str(b), style='List Bullet')

        # Skills
        skills = data.get("skills", [])
        if skills:
            doc.add_heading("TECHNICAL SKILLS", level=2)
            if isinstance(skills, list):
                doc.add_paragraph(", ".join([str(s) for s in skills if s]))
            elif isinstance(skills, dict):
                for k, v in skills.items():
                    if isinstance(v, list):
                        p = doc.add_paragraph()
                        r_k = p.add_run(f"{k.title()}: ")
                        r_k.bold = True
                        p.add_run(", ".join([str(s) for s in v]))

        # Education
        education = data.get("education", [])
        if education:
            doc.add_heading("EDUCATION", level=2)
            for edu in education:
                if isinstance(edu, dict):
                    doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('institution', '')}")

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_data = buffer.getvalue()
        buffer.close()

        return docx_data

    def _normalize_resume_data(self, resume_data: Any) -> Dict[str, Any]:
        """Convert objects or dicts to standardized resume dictionary."""
        if hasattr(resume_data, "to_dict"):
            return resume_data.to_dict()
        elif hasattr(resume_data, "dict"):
            return resume_data.dict()
        elif hasattr(resume_data, "model_dump"):
            return resume_data.model_dump()
        elif isinstance(resume_data, dict):
            return resume_data
        return {}

    def _build_styles(self, template_name: str) -> Dict[str, Any]:
        """Build ReportLab styles and color palette based on template name."""
        base_styles = getSampleStyleSheet()

        font_family = "Helvetica"
        accent_color = colors.HexColor("#1E293B")  # Slate default
        accent_hex = "#1E293B"

        if template_name == "modern":
            accent_color = colors.HexColor("#1D4ED8")  # Royal Blue
            accent_hex = "#1D4ED8"
        elif template_name == "minimal":
            accent_color = colors.HexColor("#334155")  # Dark Neutral
            accent_hex = "#334155"
        else:  # ats_classic
            accent_color = colors.HexColor("#1E1E1E")  # Pure ATS Black
            accent_hex = "#1E1E1E"

        title_style = ParagraphStyle(
            "HeaderTitle",
            parent=base_styles["Normal"],
            fontName=f"{font_family}-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER if template_name != "minimal" else TA_LEFT,
            textColor=accent_color
        )

        contact_style = ParagraphStyle(
            "HeaderContact",
            parent=base_styles["Normal"],
            fontName=font_family,
            fontSize=9,
            leading=12,
            alignment=TA_CENTER if template_name != "minimal" else TA_LEFT,
            textColor=colors.HexColor("#475569")
        )

        section_heading = ParagraphStyle(
            "SectionHeading",
            parent=base_styles["Normal"],
            fontName=f"{font_family}-Bold",
            fontSize=11,
            leading=14,
            textColor=accent_color,
            spaceAfter=4
        )

        body_style = ParagraphStyle(
            "CustomBodyText",
            parent=base_styles["Normal"],
            fontName=font_family,
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#1E293B")
        )

        summary_style = ParagraphStyle(
            "CustomSummary",
            parent=body_style,
            alignment=TA_JUSTIFY
        )

        item_header = ParagraphStyle(
            "ItemHeader",
            parent=base_styles["Normal"],
            fontName=f"{font_family}-Bold",
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#0F172A")
        )

        item_meta = ParagraphStyle(
            "ItemMeta",
            parent=base_styles["Normal"],
            fontName=f"{font_family}-Oblique",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#64748B")
        )

        bullet_style = ParagraphStyle(
            "CustomBullet",
            parent=body_style,
            leftIndent=12,
            firstLineIndent=-8,
            spaceAfter=2
        )

        return {
            "font_family": font_family,
            "accent_color": accent_color,
            "accent_hex": accent_hex,
            "Title": title_style,
            "Contact": contact_style,
            "SectionHeading": section_heading,
            "BodyText": body_style,
            "Summary": summary_style,
            "ItemHeader": item_header,
            "ItemMeta": item_meta,
            "Bullet": bullet_style
        }

    def _build_header(self, data: Dict[str, Any], template: str, styles: Dict[str, Any], story: List[Any]):
        """Render header with candidate name and contact information."""
        contact = data.get("contact", {}) or data.get("contact_info", {})
        name = contact.get("name") or data.get("name", "APPLICANT NAME")

        story.append(Paragraph(f"<b>{name.upper()}</b>", styles["Title"]))
        story.append(Spacer(1, 4))

        contact_parts = []
        for key in ["email", "phone", "location", "linkedin", "github"]:
            val = contact.get(key)
            if val:
                contact_parts.append(str(val))

        if contact_parts:
            contact_str = "  |  ".join(contact_parts)
            story.append(Paragraph(contact_str, styles["Contact"]))

        story.append(Spacer(1, 10))

    def _add_section_header(self, title: str, template: str, styles: Dict[str, Any], story: List[Any]):
        """Add styled section header with clean divider line."""
        story.append(Paragraph(f"<b>{title}</b>", styles["SectionHeading"]))
        story.append(HRFlowable(
            width="100%",
            thickness=1,
            color=styles["accent_color"],
            spaceBefore=1,
            spaceAfter=6
        ))


def generate_pdf(resume_data: Dict[str, Any], template_name: str = "ats_classic") -> bytes:
    """Convenience function to generate PDF bytes."""
    generator = PDFGenerator()
    return generator.generate_pdf(resume_data, template_name)


def generate_docx(resume_data: Dict[str, Any], template_name: str = "ats_classic") -> bytes:
    """Convenience function to generate DOCX bytes."""
    generator = PDFGenerator()
    return generator.generate_docx(resume_data, template_name)
